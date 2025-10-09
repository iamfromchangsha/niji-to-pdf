import tkinter as tk
from tkinter import messagebox, ttk
import requests
import json
import time
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
import os


def login(email, password):
    session = requests.Session()
    url = 'https://nideriji.cn/api/login/'
    data = {
        'email': email,
        'password': password,
        'csrfmiddlewaretoken': ''
    }
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0',
        'Referer': 'https://nideriji.cn/w/login'
    }
    resp = session.post(url, data=data, headers=headers)
    return resp.json()


def script(token):
    session = requests.Session()
    url = 'https://nideriji.cn/api/v2/sync/'
    headers = {
        'Auth': 'token ' + token,
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0'
    }
    data = {
        'user_config_ts': '0',
        'diaries_ts': '0',
        'readmark_ts': '0',
        'images_ts': '0'
    }
    resp = session.post(url, data=data, headers=headers)
    return resp.json()


def pin(userid, token, id):
    session = requests.Session()
    url = f'https://nideriji.cn/api/diary/all_by_ids/{userid}/'
    headers = {
        'Auth': 'token ' + token,
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0'
    }
    data = {
        'diary_ids': id
    }
    resp = session.post(url, data=data, headers=headers)
    return resp.json()


def set_chinese_font(doc, font_name='微软雅黑'):
    style = doc.styles['Normal']
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    style.font.size = Pt(12)


def chaseimg(content):
    pattern = r'\[图\d+\]'
    matches = re.findall(pattern, content)
    return matches


def get_img(id, token, userid):
    # 创建img文件夹（如果不存在）
    img_folder = 'img'
    if not os.path.exists(img_folder):
        os.makedirs(img_folder)
        
    url = f'https://f.nideriji.cn/api/image/{userid}/{id}/'
    headers = {
        'Auth': 'token ' + token,
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0'
    }
    
    # 添加重试机制
    max_retries = 3
    for attempt in range(max_retries):
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            # 保存图片到img文件夹
            img_path = os.path.join(img_folder, f"{id}.jpg")
            with open(img_path, 'wb') as f:
                f.write(response.content)
                print('✅ 图片已保存为 '+img_path)
            # 返回文件路径
            return img_path
        else:
            print(f'⚠️ 图片{id}下载失败，状态码: {response.status_code}，重试 {attempt + 1}/{max_retries}')
            time.sleep(1)  # 等待1秒后重试
    
    print(f'❌ 图片{id}下载失败，已达到最大重试次数')
    return None


def export_diaries(email, password):
    try:
        # 登录
        logins = login(email, password)
        if 'token' not in logins:
            messagebox.showerror("登录失败", "请检查邮箱或密码是否正确。")
            return

        token = logins.get('token')
        userid = logins.get('user_config', {}).get('userid')
        if not userid:
            messagebox.showerror("错误", "无法获取用户ID，请重试。")
            return

        # 获取日记列表
        shujv = script(token).get('diaries', [])
        if not shujv:
            messagebox.showinfo("提示", "没有找到日记数据。")
            return

        # 创建进度窗口
        progress_window = tk.Toplevel()
        progress_window.title("导出进度")
        progress_window.geometry("300x100")
        progress_window.resizable(False, False)
        
        progress_label = tk.Label(progress_window, text="正在导出日记...")
        progress_label.pack(pady=5)
        
        progress_var = tk.DoubleVar()
        progress_bar = ttk.Progressbar(progress_window, variable=progress_var, maximum=100)
        progress_bar.pack(padx=20, pady=10, fill=tk.X)
        
        progress_percentage = tk.Label(progress_window, text="0%")
        progress_percentage.pack()
        
        # 更新进度窗口
        progress_window.update()

        # 创建 Word 文档
        doc = Document()
        set_chinese_font(doc)

        total_diaries = len(shujv)
        for index, diary_info in enumerate(shujv):
            diary_id = diary_info.get('id')
            diary = pin(userid, token, diary_id)
            diary_data = diary.get('diaries', [{}])[0]
            content = diary_data.get('content', '')
            timets = diary_data.get('ts', 0)

            # 格式化时间
            l_time = time.localtime(timets)
            formatted_time = time.strftime('%Y-%m-%d %H:%M:%S', l_time)

            # 提取图片标签
            img_ids = chaseimg(content)
            img_files = {}
            for img_tag in img_ids:
                match = re.search(r'\[图(\d+)\]', img_tag)
                if match:
                    img_id = match.group(1)
                    img_file = get_img(img_id, token, userid)
                    img_files[img_tag] = img_file

            # 添加标题
            doc.add_heading(formatted_time, level=1)

            # 插入内容和图片
            if img_files:
                parts = re.split(r'(\[图\d+\])', content)
                for part in parts:
                    if part in img_files:
                        img_path = img_files[part]
                        if os.path.exists(img_path):
                            doc.add_picture(img_path, width=Inches(4.0))
                        else:
                            doc.add_paragraph(part)
                    elif part.strip():
                        doc.add_paragraph(part)
            else:
                doc.add_paragraph(content)
            
            # 更新进度
            progress_percent = (index + 1) / total_diaries * 100
            progress_var.set(progress_percent)
            progress_percentage.config(text=f"{int(progress_percent)}%")
            progress_window.update_idletasks()
            progress_window.update()

        # 保存文档
        output_file = 'nideriji_diaries.docx'
        doc.save(output_file)
        
        # 关闭进度窗口
        progress_window.destroy()
        
        messagebox.showinfo("成功", f"Word 文档已保存为 {output_file}")

    except Exception as e:
        messagebox.showerror("错误", f"发生异常：{str(e)}")
        # 确保在出现异常时进度窗口也被关闭
        try:
            progress_window.destroy()
        except:
            pass


def on_submit():
    email = entry_email.get().strip()
    password = entry_password.get().strip()
    if not email or not password:
        messagebox.showwarning("输入错误", "请输入邮箱和密码。")
        return
    export_diaries(email, password)


# 创建主窗口
root = tk.Tk()
root.title("你的日记导出工具")
root.geometry("400x200")

# 邮箱输入
tk.Label(root, text="邮箱：").pack(pady=(20, 5))
entry_email = tk.Entry(root, width=40)
entry_email.pack()

# 密码输入
tk.Label(root, text="密码：").pack(pady=(10, 5))
entry_password = tk.Entry(root, width=40, show="*")
entry_password.pack()

# 提交按钮
tk.Button(root, text="导出日记到 Word", command=on_submit, width=20, height=2).pack(pady=20)

# 运行 GUI
root.mainloop()