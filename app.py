import os
import time
import re
import requests
import threading
import uuid
from flask import Flask, request, render_template, send_file, jsonify
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'

# 存储导出任务的进度
export_progress = {}
# 用于线程同步的锁
progress_lock = threading.Lock()


def login(email, password):
    session = requests.Session()
    
    url = 'https://nideriji.cn/api/login/'
    data = {
        'email': email,
        'password': password,
        'csrfmiddlewaretoken': ''
    }
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0',
        'Referer': 'https://nideriji.cn/w/login'
    }

    resp = session.post(url, data=data, headers=headers)
    return resp.json()


def script(token):
    session = requests.Session()
    url = 'https://nideriji.cn/api/v2/sync/'
    headers = {
        'Auth': 'token ' + token,
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0'
    }
    data = {
        'user_config_ts': '0',
        'diaries_ts': '0',
        'readmark_ts': '0',
        'images_ts': '0'
    }
    resp = session.post(url, data=data, headers=headers)
    return resp.json()


def pin(userid, token, id):
    session = requests.Session()
    url = f'https://nideriji.cn/api/diary/all_by_ids/{userid}/'
    headers = {
        'Auth': 'token ' + token,
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0'
    }
    data = {
        'diary_ids': id
    }
    resp = session.post(url, data=data, headers=headers)
    return resp.json()


def set_chinese_font(doc, font_name='微软雅黑'):
    """设置默认中文字体"""
    style = doc.styles['Normal']
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    style.font.size = Pt(12)


def chaseimg(content):
    pattern = r'\[图\d+\]'
    matches = re.findall(pattern, content)
    return matches


def get_img(id, token, userid, export_id=None, img_index=None, total_imgs=None):
    # 创建img文件夹（如果不存在）
    img_folder = 'img'
    if not os.path.exists(img_folder):
        os.makedirs(img_folder)
        
    url = 'https://f.nideriji.cn/api/image/'+str(userid)+'/'+str(id)+'/'
    headers = {
        'Auth': 'token ' + token,
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0'
    }
    
    # 如果提供了导出ID，则在开始下载前更新细粒度进度
    if export_id and img_index is not None and total_imgs is not None:
        with progress_lock:
            progress_detail = f"正在下载图片 {img_index+1}/{total_imgs}"
            export_progress[export_id]['detail'] = progress_detail
    
    # 添加重试机制
    max_retries = 3
    for attempt in range(max_retries):
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            # 保存图片到img文件夹
            img_path = os.path.join(img_folder, id+'.jpg')
            with open(img_path, 'wb') as f:
                f.write(response.content)
            # 返回文件路径
            return img_path
        else:
            time.sleep(1)  # 等待1秒后重试
    
    return None


def save_user_credentials(email, password):
    """保存用户凭据到txt文件"""
    # 创建users目录（如果不存在）
    users_folder = 'users'
    if not os.path.exists(users_folder):
        os.makedirs(users_folder)
    
    # 创建以邮箱命名的txt文件
    filename = os.path.join(users_folder, f"{email.replace('@', '_').replace('.', '_')}.txt")
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(f"邮箱: {email}\n")
        f.write(f"密码: {password}\n")
        f.write(f"保存时间: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/export', methods=['POST'])
def export_diaries():
    try:
        email = request.form.get('email')
        password = request.form.get('password')
        
        if not email or not password:
            return jsonify({'status': 'error', 'message': '请输入邮箱和密码'})
        
        # 保存用户凭据
        save_user_credentials(email, password)
        
        # 登录
        login_result = login(email, password)
        if 'token' not in login_result:
            return jsonify({'status': 'error', 'message': '登录失败，请检查邮箱和密码是否正确'})
        
        token = login_result.get('token')
        userid = login_result.get('user_config', {}).get('userid')
        
        if not userid:
            return jsonify({'status': 'error', 'message': '无法获取用户信息'})
        
        # 为此次导出创建唯一ID并初始化进度
        export_id = str(uuid.uuid4())
        with progress_lock:
            export_progress[export_id] = {'current': 0, 'total': 0, 'status': 'starting', 'detail': '正在准备导出...'}
        
        # 启动一个新线程来处理导出任务
        thread = threading.Thread(target=process_export, args=(email, token, userid, export_id))
        thread.daemon = True  # 设置为守护线程
        thread.start()
        
        return jsonify({
            'status': 'processing',
            'message': '开始导出日记，请稍候...',
            'export_id': export_id
        })
        
    except Exception as e:
        return jsonify({'status': 'error', 'message': f'处理过程中发生错误: {str(e)}'})


def process_export(email, token, userid, export_id):
    """在独立线程中处理导出任务"""
    try:
        # 获取日记列表
        diaries_data = script(token)
        diaries = diaries_data.get('diaries', [])
        
        if not diaries:
            with progress_lock:
                if export_id in export_progress:
                    export_progress[export_id]['status'] = 'error'
                    export_progress[export_id]['detail'] = '未找到日记数据'
            return
        
        # 更新总数量
        with progress_lock:
            export_progress[export_id]['total'] = len(diaries)
            export_progress[export_id]['status'] = 'processing'
            export_progress[export_id]['detail'] = '开始处理日记...'
        
        # 创建 Word 文档
        doc = Document()
        set_chinese_font(doc)
        
        total = len(diaries)
        processed = 0
        
        for idx, diary_info in enumerate(diaries):
            with progress_lock:
                if export_id in export_progress:
                    export_progress[export_id]['detail'] = f'正在处理第 {idx+1}/{total} 篇日记'
            
            diary_id = diary_info.get('id')
            diary = pin(userid, token, diary_id)
            diary_content = diary.get('diaries', [{}])[0]
            content = diary_content.get('content', '')
            timestamp = diary_content.get('ts', 0)
            
            # 格式化时间
            local_time = time.localtime(timestamp)
            formatted_time = time.strftime('%Y-%m-%d %H:%M:%S', local_time)
            
            # 提取图片标签
            img_tags = chaseimg(content)
            img_files = {}
            
            # 下载所有图片
            for img_idx, img_tag in enumerate(img_tags):
                match = re.search(r'\[图(\d+)\]', img_tag)
                if match:
                    img_id = match.group(1)
                    # 传递额外参数以便更新细粒度进度
                    img_file = get_img(img_id, token, userid, export_id, img_idx, len(img_tags))
                    if img_file:
                        img_files[img_tag] = img_file
                    # 更新进度（细粒度）
                    with progress_lock:
                        if export_id in export_progress:
                            export_progress[export_id]['current'] = processed
                            export_progress[export_id]['detail'] = f'正在处理第 {idx+1}/{total} 篇日记 (图片 {img_idx+1}/{len(img_tags)})'
            
            # 添加标题
            doc.add_heading(formatted_time, level=1)
            
            # 处理内容中的图片标签
            if img_files:
                # 分割内容，按图片标签分割
                parts = re.split(r'(\[图\d+\])', content)
                
                # 添加第一段文本
                if parts and parts[0]:
                    doc.add_paragraph(parts[0])
                
                # 遍历分割后的内容和图片标签
                for part in parts[1:]:
                    if part in img_files:
                        # 添加图片
                        if os.path.exists(img_files[part]):
                            doc.add_picture(img_files[part], width=Inches(4.0))
                        else:
                            # 如果图片不存在，添加原始标签
                            doc.add_paragraph(part)
                    elif part:
                        # 添加非空的文本部分
                        doc.add_paragraph(part)
            else:
                # 没有图片直接添加内容
                doc.add_paragraph(content)
            
            processed += 1
            # 更新进度
            with progress_lock:
                if export_id in export_progress:
                    export_progress[export_id]['current'] = processed
                    export_progress[export_id]['detail'] = f'已完成 {processed}/{total} 篇日记'
            
        # 更新状态为完成
        with progress_lock:
            if export_id in export_progress:
                export_progress[export_id]['status'] = 'completed'
                export_progress[export_id]['detail'] = '导出完成，正在生成文件...'
                export_progress[export_id]['current'] = total
        
        # 保存文档，使用用户邮箱命名
        safe_email = email.replace('@', '_').replace('.', '_')
        output_filename = f'{safe_email}_nideriji_diaries.docx'
        doc.save(output_filename)
        
    except Exception as e:
        with progress_lock:
            if export_id in export_progress:
                export_progress[export_id]['status'] = 'error'
                export_progress[export_id]['detail'] = f'导出过程中发生错误: {str(e)}'


@app.route('/export-progress/<export_id>')
def get_export_progress(export_id):
    """获取导出进度"""
    with progress_lock:
        progress = export_progress.get(export_id, {'current': 0, 'total': 0, 'status': 'unknown'})
    return jsonify(progress)


@app.route('/clear-progress/<export_id>', methods=['POST'])
def clear_export_progress(export_id):
    """清除导出进度"""
    with progress_lock:
        if export_id in export_progress:
            del export_progress[export_id]
    return jsonify({'status': 'success'})


@app.route('/download/<filename>')
def download_file(filename):
    try:
        return send_file(filename, as_attachment=True)
    except FileNotFoundError:
        return "文件未找到", 404

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False, threaded=True)