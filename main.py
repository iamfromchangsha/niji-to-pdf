import requests
import json
import time
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
import os

def login(email, password):
    session = requests.Session()
    
    url = 'https://nideriji.cn/api/login/'
    data = {
        'email': email,
        'password': password,
        'csrfmiddlewaretoken': ''  # 尝试空字符串
    }
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0',
        'Referer': 'https://nideriji.cn/w/login'
    }

    resp = session.post(url, data=data, headers=headers)
    return resp.json()

def script(token):
    session = requests.Session()
    url = 'https://nideriji.cn/api/v2/sync/'
    headers = {
        'Auth': 'token ' + token,
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0'
    }
    data = {
        'user_config_ts': '0',
        'diaries_ts': '0',
        'readmark_ts': '0',
        'images_ts': '0'
    }
    resp = session.post(url, data=data, headers=headers)
    return resp.json()

def pin(userid, token, id):
    session = requests.Session()
    url = f'https://nideriji.cn/api/diary/all_by_ids/{userid}/'
    headers = {
        'Auth': 'token ' + token,
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0'
    }
    data = {
        'diary_ids': id
    }
    resp = session.post(url, data=data, headers=headers)
    return resp.json()

def set_chinese_font(doc, font_name='微软雅黑'):
    """设置默认中文字体"""
    style = doc.styles['Normal']
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    style.font.size = Pt(12)

def chaseimg(content):
    pattern = r'\[图\d+\]'
    matches = re.findall(pattern, content)
    return matches

def get_img(id,token,userid):
    url = 'https://f.nideriji.cn/api/image/'+str(userid)+'/'+str(id)+'/'
    headers = {
        'Auth': 'token ' + token,
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/141.0.0.0 Safari/537.36 Edg/141.0.0.0'
    }
    response = requests.get(url, headers=headers)
    with open(id+'.jpg', 'wb') as f:
        f.write(response.content)
        print('✅ 图片已保存为 '+id+'.jpg')
    # 返回文件名
    return id+'.jpg'

def main():
    logins = login('', '')
    token = logins.get('token')
    userid = logins.get('user_config').get('userid')
    print(userid)
    
    shujv = script(token).get('diaries')
    print(len(shujv))

    # 创建 Word 文档
    doc = Document()
    set_chinese_font(doc)  # 设置中文字体

    for i in shujv:
        diary = pin(userid, token, i.get('id'))
        content = diary.get('diaries')[0].get('content')
        img_ids = chaseimg(content)
        img_files = {}  # 存储图片ID和文件路径的映射
        
        # 下载所有图片
        for img_tag in img_ids:
            pattern = r'\[图(\d+)\]'
            numbers = re.findall(pattern, img_tag)
            if numbers:
                img_id = numbers[0]
                # 下载图片并保存文件路径
                img_file = get_img(img_id, token, userid)
                img_files[img_tag] = img_file

        timets = diary.get('diaries')[0].get('ts')
        l_time = time.localtime(timets)
        formatted_time = time.strftime('%Y-%m-%d %H:%M:%S', l_time)

        # 添加标题（一级标题）
        doc.add_heading(formatted_time, level=1)
        
        # 处理内容中的图片标签
        if img_files:
            # 分割内容，按图片标签分割
            parts = re.split(r'(\[图\d+\])', content)            
            # 添加第一段文本（图片标签前的内容）
            if parts and parts[0]:
                doc.add_paragraph(parts[0])
            
            # 遍历分割后的内容和图片标签
            for j in range(1, len(parts)):
                part = parts[j]
                if part in img_files:
                    # 添加图片
                    if os.path.exists(img_files[part]):
                        doc.add_picture(img_files[part], width=Inches(4.0))
                    else:
                        # 如果图片不存在，添加原始标签
                        doc.add_paragraph(part)
                elif part:  # 添加非空的文本部分
                    doc.add_paragraph(part)
        else:
            # 没有图片直接添加内容
            doc.add_paragraph(content)

        print(formatted_time)
        print(content)

    # 保存文档
    doc.save('nideriji_diaries.docx')
    print('✅ Word 文档已保存为 nideriji_diaries.docx')

if __name__ == '__main__':
    main()